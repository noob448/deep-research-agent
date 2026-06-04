"""
将 Markdown 研究报告转换为 .docx 文件。

纯 Python 实现（python-docx），不依赖 pandoc。
格式化工作是确定性任务——留在 agent 循环之外更可靠。
"""

import re
from pathlib import Path

from .config import WORKSPACE_DIR, REPORT_FILENAME, OUTPUT_DOCX_FILENAME

# python-docx 是可选的——如果没装就跳过 .docx 生成
try:
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False


def convert_report(md_path: Path | None = None, docx_path: Path | None = None) -> Path:
    """将 Markdown 报告转换为 .docx 文件。

    Args:
        md_path: Markdown 报告路径，默认为 workspace/report.md
        docx_path: 输出 .docx 路径，默认为 workspace/report.docx

    Returns:
        生成的 .docx 文件路径

    Raises:
        ImportError: 如果 python-docx 未安装且环境有冲突的 docx 包
    """
    if not HAS_DOCX:
        raise ImportError(
            "python-docx 未正确安装。请运行: pip uninstall docx -y && pip install python-docx"
        )

    md_path = md_path or WORKSPACE_DIR / REPORT_FILENAME
    docx_path = docx_path or WORKSPACE_DIR / OUTPUT_DOCX_FILENAME

    if not md_path.exists():
        raise FileNotFoundError(f"报告文件不存在: {md_path}")

    markdown = md_path.read_text(encoding="utf-8")
    doc = Document()

    # 设置默认字体
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Microsoft YaHei"
    font.size = Pt(11)

    # 逐行解析
    lines = markdown.split("\n")
    i = 0
    while i < len(lines):
        line = lines[i]

        # 空行
        if not line.strip():
            i += 1
            continue

        # 标题
        if line.startswith("# "):
            _add_heading(doc, line[2:].strip(), level=1)
        elif line.startswith("## "):
            _add_heading(doc, line[3:].strip(), level=2)
        elif line.startswith("### "):
            _add_heading(doc, line[4:].strip(), level=3)
        elif line.startswith("#### "):
            _add_heading(doc, line[5:].strip(), level=4)

        # 无序列表
        elif line.strip().startswith("- ") or line.strip().startswith("* "):
            text = line.strip()[2:]
            _add_list_item(doc, text)

        # 有序列表
        elif re.match(r"^\d+\.\s", line.strip()):
            text = re.sub(r"^\d+\.\s", "", line.strip())
            _add_list_item(doc, text)

        # 任务列表
        elif re.match(r"^- \[[ x]\]\s", line.strip()):
            text = re.sub(r"^- \[[ x]\]\s", "", line.strip())
            checked = "[x]" in line[:6]
            _add_task_item(doc, text, checked)

        # 代码块
        elif line.startswith("```"):
            code_lines = []
            i += 1
            while i < len(lines) and not lines[i].startswith("```"):
                code_lines.append(lines[i])
                i += 1
            _add_code_block(doc, "\n".join(code_lines))

        # 普通段落
        else:
            _add_paragraph(doc, line)

        i += 1

    doc.save(str(docx_path))
    return docx_path


def _add_heading(doc: Document, text: str, level: int):
    """添加标题。"""
    heading = doc.add_heading(text, level=level)
    # 标题字体调整
    for run in heading.runs:
        run.font.name = "Microsoft YaHei"
        if level == 1:
            run.font.size = Pt(18)
        elif level == 2:
            run.font.size = Pt(14)
        elif level == 3:
            run.font.size = Pt(12)


def _add_paragraph(doc: Document, text: str):
    """添加段落，处理行内格式。"""
    para = doc.add_paragraph()
    _add_formatted_runs(para, text)


def _add_list_item(doc: Document, text: str):
    """添加列表项。"""
    para = doc.add_paragraph(style="List Bullet")
    # 清除默认文本后添加格式化 runs
    para.clear()
    _add_formatted_runs(para, text)


def _add_task_item(doc: Document, text: str, checked: bool):
    """添加任务列表项。"""
    prefix = "☑ " if checked else "☐ "
    para = doc.add_paragraph()
    _add_formatted_runs(para, prefix + text)


def _add_code_block(doc: Document, code: str):
    """添加代码块。"""
    para = doc.add_paragraph()
    run = para.add_run(code)
    run.font.name = "Consolas"
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)


def _add_formatted_runs(para, text: str):
    """解析行内格式（粗体、斜体、链接）并添加到段落。"""
    # 匹配: **粗体**, *斜体*, [文本](url), 普通文本
    # 5 个捕获组: (全匹配, 粗体内容, 斜体内容, 链接文本, 链接URL)
    pattern = r"(\*\*(.+?)\*\*|\*(.+?)\*|\[(.+?)\]\((.+?)\)|[^*\[\]]+)"
    parts = re.findall(pattern, text)

    for part in parts:
        full_match, bold_text, italic_text, link_text, link_url = part

        if bold_text:
            run = para.add_run(bold_text)
            run.bold = True
        elif italic_text:
            run = para.add_run(italic_text)
            run.italic = True
        elif link_text:
            run = para.add_run(f"{link_text} ({link_url})")
            run.font.color.rgb = RGBColor(0x05, 0x63, 0xC1)
            run.underline = True
        elif full_match:
            para.add_run(full_match)
